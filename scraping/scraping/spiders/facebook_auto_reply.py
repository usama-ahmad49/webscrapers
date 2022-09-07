import time
import aspose.words as aw
import PyPDF2
from docx2pdf import convert
import docx
import os
from selenium import webdriver
from selenium.webdriver import ActionChains
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

if __name__ == '__main__':
    cwd = os.getcwd()
    # open document contaning messgage for auto reply
    # we use docx library for this purpose
    convert("Auto response for Facebook Marketplace.docx",cwd+'\\'"Auto response for Facebook Marketplace.pdf")
    # doc = aw.Document("Auto response for Facebook Marketplace.docx")
    #
    # # Save as PDF
    # doc.save("Auto response for Facebook Marketplace.pdf")
    # try:
    #     doc = docx.Document('Auto response for Facebook Marketplace.docx')
    #     AutomatedMessage = ''  # variable contains message that will be sent as reply
    #     fullText = []
    #     for para in doc.paragraphs:
    #         fullText.append(para.text)
    #     AutomatedMessage = '\r'.join(fullText)
    # except IOError:
    #     # In case of any problem opening file system will print error and close itself.
    #     print('error opening file')
    #     exit()

    # open file containg facebook account information
    # Informaton in file must be stored in a perticular pattern to avoid any errors
    # all account information must be in following pattern:
    # useremail>password
    # for adding more then one account in file:
    # add an account info in above pattern then press enter one time then add another account info in above pattern

    file = open("Auto response for Facebook Marketplace.pdf", 'rb')
    pdfReader = PyPDF2.PdfFileReader(file)
    pageObj = pdfReader.getPage(0)
    AutomatedMessage = pageObj.extractText()
    with open("Facebook_Accounts_file.txt", 'r', encoding='utf-8') as FAF:
        DATA = FAF.read().split('\n')

    options = Options()
    options.add_argument('--disable-notifications')  # to disable any unwanted notification popups or alerts
    driver = webdriver.Chrome(options=options)  # open browser
    driver.maximize_window()  # maximize window
    driver.get('https://www.facebook.com/')  # go to facebook.om login page
    for data in DATA:
        if data == '':
            continue
        username = data.split('>')[0]
        password = data.split('>')[1]

        WebDriverWait(driver, 30).until(
            EC.element_to_be_clickable((By.ID, 'email')))  # wait for browser to load elements that we need to click
        time.sleep(3)

        driver.find_element(By.ID, 'email').send_keys(username)  # enter username
        time.sleep(0.5)
        driver.find_element(By.ID, 'pass').send_keys(password)  # enter password
        time.sleep(0.5)
        driver.find_element(By.CSS_SELECTOR, "button[name='login']").click()  # clink login button
        WebDriverWait(driver, 30).until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'input[type="search"]')))
        time.sleep(3)

        driver.get('https://www.facebook.com/marketplace/inbox')  # go to marketplace inbox after logging in
        time.sleep(3)
        WebDriverWait(driver, 30).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, 'div[aria-label="Collection of Marketplace items"]')))
        driver.find_element(By.CSS_SELECTOR, 'div[aria-label="Collection of Marketplace items"]')
        time.sleep(3)
        # get all messages from inbox and loop over them one by one to reply to them individually
        for messages in driver.find_elements(By.CSS_SELECTOR,
                                             'div[aria-label="Collection of Marketplace items"] div[data-visualcompletion="ignore-dynamic"]')[
                        1:]:
            a = ActionChains(driver)
            a.move_to_element(messages).perform()
            time.sleep(1)
            messages.click()  # open message box
            time.sleep(3)
            WebDriverWait(driver, 30).until(EC.element_to_be_clickable((By.CSS_SELECTOR, 'div[role="textbox"]')))
            # click text box in message window
            driver.find_element(By.CSS_SELECTOR, 'div[role="textbox"]').click()
            time.sleep(1)
            # clear text box in message window to make sure no extra keystrocks sent to user
            driver.find_element(By.CSS_SELECTOR, 'div[role="textbox"]').clear()
            time.sleep(1)
            # paste coppied automated message into the text box to be sent
            driver.find_element(By.CSS_SELECTOR, 'div[role="textbox"]').send_keys(AutomatedMessage.replace('\n', '\r'))
            time.sleep(1)
            # press enter to send message
            driver.find_element(By.CSS_SELECTOR, 'div[role="textbox"]').send_keys(Keys.ENTER)
            time.sleep(1)
            # now message sent! close the message window
            driver.find_element(By.CSS_SELECTOR, 'div[aria-label="Close chat"]').click()
            time.sleep(3)
        # click on the profile element to show logout button
        try:
            driver.find_element(By.CSS_SELECTOR, 'div[aria-label="Your profile"]').click()
        except:
            driver.find_element(By.CSS_SELECTOR, 'svg[aria-label="Your profile"]').click()
        time.sleep(1)
        # click logout button
        driver.find_element(By.CSS_SELECTOR,
                            'div[data-visualcompletion="ignore-dynamic"][role="listitem"][data-nocookies="true"]').click()

    # all accounts served now close browser and end program
    driver.quit()
